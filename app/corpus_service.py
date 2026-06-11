# app/corpus_service.py
"""Corpus/RAG substrate: ingest, chunk, embed, and cosine-search career docs.

Embeddings are stored as float32 BLOBs and searched by brute-force numpy cosine
behind `search()` (swappable to sqlite-vec later). The embedder is injectable so
the default test suite runs offline; the default wraps OpenAI text-embedding-3-small.
"""

from __future__ import annotations

import hashlib
import io
from collections.abc import Callable
from dataclasses import dataclass

import numpy as np
from sqlmodel import Session, delete, select

from app import settings_service
from app.config import get_config
from app.models import (
    Chunk,
    Document,
    DocumentMediaType,
    DocumentSource,
)

# An embedder maps a batch of texts to a batch of vectors.
Embedder = Callable[[list[str]], list[list[float]]]


def extract_text(*, data: bytes, media_type: DocumentMediaType) -> str:
    """Extract plain text from raw bytes by media type. Raises ValueError if empty."""
    if media_type in (DocumentMediaType.txt, DocumentMediaType.md):
        text = data.decode("utf-8", errors="replace")
    elif media_type == DocumentMediaType.pdf:
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(data))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    elif media_type == DocumentMediaType.docx:
        import docx

        document = docx.Document(io.BytesIO(data))
        text = "\n".join(p.text for p in document.paragraphs)
    else:  # pragma: no cover - enum is exhaustive
        raise ValueError(f"unsupported media type: {media_type}")

    text = text.strip()
    if not text:
        raise ValueError("no extractable text in document")
    return text


def chunk_text(text: str, *, size: int = 1000, overlap: int = 150) -> list[str]:
    """Split text into overlapping windows, breaking on a nearby boundary.

    Deterministic and dependency-free so it is trivially testable.
    """
    text = text.strip()
    if not text:
        return []
    chunks: list[str] = []
    n = len(text)
    start = 0
    while start < n:
        end = min(start + size, n)
        if end < n:
            window = text[start:end]
            # Prefer the most semantic boundary available past the window
            # midpoint: paragraph, then line, then sentence, then word.
            for pattern in ("\n\n", "\n", ". ", " "):
                brk = window.rfind(pattern)
                if brk > size // 2:
                    end = start + brk + 1
                    break
        piece = text[start:end].strip()
        if piece:
            chunks.append(piece)
        if end >= n:
            break
        start = max(end - overlap, start + 1)
    return chunks


def default_embedder(session: Session) -> Embedder:
    """Build an OpenAI-backed embedder using the settings-resolved key.

    Raises RuntimeError (not at import) if no key is configured.
    """
    api_key = settings_service.resolve_openai_key(session)
    if not api_key:
        raise RuntimeError("OpenAI API key is not configured (Settings or OH_OPENAI_API_KEY).")
    model = get_config().embedding_model

    def embed(texts: list[str]) -> list[list[float]]:
        import openai

        client = openai.OpenAI(api_key=api_key)
        resp = client.embeddings.create(model=model, input=texts)
        return [item.embedding for item in resp.data]

    return embed


def _to_blob(vector: list[float]) -> bytes:
    return np.asarray(vector, dtype=np.float32).tobytes()


def ingest_document(
    session: Session,
    *,
    title: str,
    source_kind: DocumentSource,
    media_type: DocumentMediaType,
    data: bytes,
    embedder: Embedder,
) -> Document:
    """Extract → hash → (dedup-replace) → chunk → embed → persist atomically."""
    raw_text = extract_text(data=data, media_type=media_type)
    content_hash = hashlib.sha256(raw_text.encode("utf-8")).hexdigest()

    # Idempotency: drop any existing document (and its chunks) with the same hash.
    existing = session.exec(
        select(Document).where(Document.content_hash == content_hash)
    ).all()
    for old in existing:
        session.exec(delete(Chunk).where(Chunk.document_id == old.id))
        session.delete(old)
    session.flush()

    pieces = chunk_text(raw_text)
    if not pieces:
        raise ValueError("document produced no chunks")
    vectors = embedder(pieces)
    if len(vectors) != len(pieces):
        raise ValueError("embedder returned wrong number of vectors")

    model = get_config().embedding_model
    doc = Document(
        title=title, source_kind=source_kind, media_type=media_type,
        raw_text=raw_text, content_hash=content_hash, char_count=len(raw_text),
    )
    session.add(doc)
    session.flush()  # assign doc.id
    for seq, (piece, vec) in enumerate(zip(pieces, vectors)):
        session.add(Chunk(
            document_id=doc.id, seq=seq, text=piece,
            embedding=_to_blob(vec), embedding_model=model,
        ))
    session.commit()
    session.refresh(doc)
    return doc
